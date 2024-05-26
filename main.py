import openpyxl
from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml import parse_xml
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from tqdm import tqdm

def set_font(document, font_name):
    # Set the font of an element and its children in a Word document.
    styles = document.styles
    default_style = styles['Normal']
    default_font = default_style.font
    default_font.name = font_name
    default_font.size = Pt(11)  # Set the desired font size

def set_table_borders(table):
    border_xml = """
        <w:tblBorders xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">
            <w:top w:val="single" w:sz="4" w:space="0" w:color="auto"/>
            <w:left w:val="single" w:sz="4" w:space="0" w:color="auto"/>
            <w:right w:val="single" w:sz="4" w:space="0" w:color="auto"/>
            <w:bottom w:val="single" w:sz="4" w:space="0" w:color="auto"/>
        </w:tblBorders>
    """
    table._element.xpath('//w:tblPr')[0].append(parse_xml(border_xml))

def prevent_table_split(table):
    # Prevent the table from being split across pages
    tbl = table._element
    tblPr = tbl.xpath("w:tblPr")
    if not tblPr:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)
    else:
        tblPr = tblPr[0]

    cantSplit = OxmlElement('w:cantSplit')
    tblPr.append(cantSplit)

    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.keep_with_next = True
                paragraph.paragraph_format.keep_together = True

def get_table_height(table):
    # Estimate the height of the table by summing the heights of the rows
    height = 0
    for row in table.rows:
        row_height = 0
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    row_height += run.font.size.pt if run.font.size else 11
        height += row_height
    return height / 72  # Convert points to inches

def create_tables_from_excel_rows(excel_file_path, sheet_name, word_file_path):
    # Load Excel workbook and select worksheet
    workbook = openpyxl.load_workbook(excel_file_path)
    worksheet = workbook[sheet_name]

    # Sort Excel rows alphabetically based on the first column, ignoring case
    sorted_rows = sorted(worksheet.iter_rows(min_row=2, values_only=True), key=lambda row: row[0].casefold() if row[0] else '')

    # Create a new Word document
    doc = Document()

    # Set the default font to Arial
    set_font(doc, 'Arial')

    # Track the current letter
    current_letter = ''

    section = doc.sections[0]
    available_height = section.page_height - section.top_margin - section.bottom_margin

    # Loop through each sorted row in Excel with progress bar
    for row in tqdm(sorted_rows, desc="Processing index"):
        # Get the first character of the entry1 value
        entry1_value = row[0] if row[0] else None
        first_letter = entry1_value[0].upper() if entry1_value else ''

        # Check if a new letter section is starting
        if first_letter != current_letter:
            current_letter = first_letter

            # Add a page break before the new letter section
            doc.add_page_break()

            # Add a header for the letter section
            header_paragraph = doc.add_paragraph()
            header_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            run = header_paragraph.add_run()
            run.text = current_letter.upper() + current_letter.lower()
            run.bold = True
            run.font.size = Pt(36)

        # Create a new table in Word with borders
        table = doc.add_table(rows=2 if row[3] else 1, cols=2)
        table.style = 'Table Grid'

        # Set border properties for the table
        set_table_borders(table)

        # Populate the table cells with the data from Excel
        entry1_cell = table.cell(0, 0)
        entry1_value = row[0] if row[0] else None  # Entry1
        if entry1_value is not None:
            entry1_cell.merge(table.cell(0, 1))
            entry1_cell.text = str(entry1_value)
            entry1_cell.paragraphs[0].runs[0].bold = True
            entry1_cell.paragraphs[0].runs[0].font.size = Pt(12)

        if row[3]:  # If there is a description
            description_cell = table.cell(1, 0)
            description_value = row[3] if row[3] else None  # Description
            if description_value is not None:
                description_cell.merge(table.cell(1, 1))
                description_cell.text = str(description_value)

        table_height = get_table_height(table)
        if table_height > available_height:
            doc.add_page_break()

        prevent_table_split(table)  # Prevent table from splitting across pages

        # Create a new table for Book and Page details
        table_details = doc.add_table(rows=1, cols=2)
        table_details.style = 'Table Grid'

        # Set border properties for the table
        set_table_borders(table_details)

        book_cell = table_details.cell(0, 0)
        book_value = row[2] if row[2] else None  # Book
        if book_value is not None:
            book_cell.text = f"Book: {str(book_value)}"
            book_cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            book_cell.paragraphs[0].paragraph_format.alignment = WD_ALIGN_VERTICAL.CENTER
            book_cell.paragraphs[0].runs[0].bold = True

        # Set border properties for the Page and Book cells
        page_cell = table_details.cell(0, 1)
        pages = row[1] if row[1] else None
        if pages is not None:
            pages = str(pages)
            if ',' in pages or '-' in pages:
                page_cell.text = f"Pages: {pages}"
            else:
                page_cell.text = f"Page: {pages}"
            page_cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            page_cell.paragraphs[0].paragraph_format.alignment = WD_ALIGN_VERTICAL.CENTER
            page_cell.paragraphs[0].runs[0].bold = True

        # Add an empty paragraph after the tables
        doc.add_paragraph()

    # Save the Word document
    doc.save(word_file_path)

    print("Your awesome index has been generated successfully!")

# Example usage
excel_file_path = "Index.xlsx"
sheet_name = "Sheet1"
word_file_path = "Awesome-Index.docx"

create_tables_from_excel_rows(excel_file_path, sheet_name, word_file_path)
